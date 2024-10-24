import os
import time

from bs4 import BeautifulSoup
from docx import Document, shared
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

class DocumentManager:
    def __init__(self, project_dir, name_project, msg_manager):
        self.project_dir = project_dir
        self.name_project = name_project
        self.msg_manager = msg_manager
        self.__doc = self.load_document()

    @property
    def doc(self):
        return self.__doc

    # Функция для сохранения HTML-кода страницы в файл
    def save_html(self, page_source, volume, chapter, project_dir):
        # Формируем имя файла, включая номер тома и главы
        html_file_path = os.path.join(project_dir, f"volume_{volume}_chapter_{chapter}.html")
        
        # Сохраняем исходный код страницы в файл
        with open(html_file_path, 'w', encoding='utf-8') as file:
            file.write(page_source)

        self.msg_manager.show_message('html_chapter_saved', chapter)

    # Загрузка уже существующего документа или создание нового
    def load_document(self):
        docx_file_path = os.path.join(self.project_dir, f"{self.name_project}.docx")

        if os.path.exists(docx_file_path):
            self.msg_manager.show_message('uploading_document')
            return Document(docx_file_path)
        else:
            self.msg_manager.show_message('document_not_found')
            return Document()

    # Закрытие документа после завершения работы
    def save_document(self):
        docx_file_path = os.path.join(self.project_dir, f"{self.name_project}.docx")
        self.__doc.save(docx_file_path)
        self.msg_manager.show_message('all_chapters_are_saved_to', docx_file_path)

    # Сохранение глав
    def save_chapters(self, chapter, doc, driver, project):

        start_time = time.time()

        # Ждем загрузки всего текста
        WebDriverWait(driver, 50).until(EC.presence_of_element_located((By.TAG_NAME, 'p')))

        # Получаем HTML страницы
        page_source = driver.page_source

        # Парсим HTML с помощью BeautifulSoup
        soup = BeautifulSoup(page_source, 'html.parser')

        # Сохраняем HTML, если проект особый
        if project.is_special:
            self.save_html(page_source, project.volume, project.chapter, project.project_dir)

        # Поиск заголовков (элементы h1) и параграфов (элементы p)
        header_h1 = soup.find_all('h1')
        header_p = soup.find_all('p')

        # Извлечение текста заголовка
        h1_texts = [h1.get_text(strip=True) for h1 in header_h1]
        chapter_title = h1_texts[0] if h1_texts else f"Глава {chapter}"

        # Добавляем заголовок главы в документ
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(chapter_title)
        title_run.bold = True  # Делаем текст жирным
        title_run.font.size = shared.Pt(14)  # Устанавливаем размер шрифта (например, 14 пт)

        # Объединяем текст параграфов
        all_text = '\n'.join([p.get_text(strip=True) for p in header_p if p.get_text(strip=True)])  
        cleaned_text = '\n'.join([line.strip() for line in all_text.splitlines() if line.strip()])  # Убираем пустые строки

        # Добавляем очищенный текст в документ
        for line in cleaned_text.splitlines():
            doc.add_paragraph(line)  # Добавление параграфа в документ
        
        doc.add_page_break()  # Разрыв страницы после каждой главы
        
        end_time = time.time()
        elapsed_time = end_time - start_time  # Вычисляем затраченное время
        self.msg_manager.show_message('time_spent', chapter, elapsed_time)
        self.msg_manager.show_message('html_chapter_saved', chapter)
