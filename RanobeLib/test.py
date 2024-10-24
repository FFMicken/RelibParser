import os

from bs4 import BeautifulSoup
from docx import Document, shared
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

class DocumentManager:
    def __init__(self, project_dir, name_project, msg_manager):
        self.project_dir = project_dir
        self.name_project = name_project
        self.msg_manager = msg_manager
        self.__doc = self.load_document()

    @property
    def doc(self):
        return self.__doc

    # Загрузка уже существующего документа или создание нового
    def load_document(self):
        docx_file_path = os.path.join(self.project_dir, f"{self.name_project}.docx")

        if os.path.exists(docx_file_path):
            self.msg_manager.show_message('uploading_document')
            return Document(docx_file_path)
        else:
            self.msg_manager.show_message('document_not_found')
            return Document()

    # Закрытие документа после завершения работы
    def save_document(self):
        docx_file_path = os.path.join(self.project_dir, f"{self.name_project}.docx")
        self.__doc.save(docx_file_path)
        self.msg_manager.show_message('all_chapters_are_saved_to', docx_file_path)

    # Основной метод для сохранения главы
    def save_chapters(self, chapter, driver, project):
        WebDriverWait(driver, 50).until(EC.presence_of_element_located((By.TAG_NAME, 'body')))
        
        # Извлечение текста и заголовка через bs4
        soup = BeautifulSoup(driver.page_source, 'html.parser')

        # Сохранение HTML для специального проекта
        if project.is_special:
            self.save_html(driver.page_source, project.volume, project.chapter, project.project_dir)

        # Получение и сохранение заголовка главы
        chapter_title = self.extract_chapter_title(soup, chapter)
        self.add_title_to_document(chapter_title)

        # Извлечение и добавление текста главы
        chapter_text = self.extract_chapter_text(soup)
        self.add_text_to_document(chapter_text)

        self.__doc.add_page_break()
        self.msg_manager.show_message('html_chapter_saved', chapter)

    # Извлечение заголовка главы
    def extract_chapter_title(self, soup, chapter):
        header_h1 = soup.find('h1')
        return header_h1.get_text(strip=True) if header_h1 else f"Глава {chapter}"

    # Извлечение текста главы
    def extract_chapter_text(self, soup):
        paragraphs = soup.find_all('p')
        return '\n'.join(p.get_text(strip=True) for p in paragraphs if p.get_text(strip=True))

    # Добавление заголовка в документ
    def add_title_to_document(self, chapter_title):
        title_paragraph = self.__doc.add_paragraph()
        title_run = title_paragraph.add_run(chapter_title)
        title_run.bold = True
        title_run.font.size = shared.Pt(14)

    # Добавление текста главы в документ
    def add_text_to_document(self, chapter_text):
        for line in chapter_text.splitlines():
            self.__doc.add_paragraph(line)

    # Функция для сохранения HTML-кода страницы в файл
    def save_html(self, page_source, volume, chapter, project_dir):
        html_file_path = os.path.join(project_dir, f"volume_{volume}_chapter_{chapter}.html")
        with open(html_file_path, 'w', encoding='utf-8') as file:
            file.write(page_source)
        self.msg_manager.show_message('html_chapter_saved', chapter)
