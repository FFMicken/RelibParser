import os
from docx import Document, shared

class DocxManager:
    def __init__(self, project_dir, name_project, msg_manager):
        self.project_dir = project_dir
        self.name_project = name_project
        self.msg_manager = msg_manager
        self.chapter_buffer = []
        self.buffer_limit = 50
        self.docx_file_path = os.path.join(self.project_dir, f"{self.name_project}.docx")

        if not os.path.exists(self.docx_file_path):
            self.doc = Document()
        else:
            self.doc = Document(self.docx_file_path)

    def save_main_page_info(self, data):
        try:
            self.doc.add_heading('Информация о книге', level=1)
            
            for key, value in data.items():
                if value:
                    if isinstance(value, list):
                        value = ', '.join(value) 
                    self.doc.add_paragraph(f"{key}: {value}")
            
            self.doc.add_page_break()
            self.doc.save(self.docx_file_path)

        except Exception as e:
            self.msg_manager.show_message('there_was_an_error', e)

    def flush_buffer(self):
            try:
                for chapter_title, chapter_text in self.chapter_buffer:
                    title_paragraph = self.doc.add_paragraph()
                    title_run = title_paragraph.add_run(chapter_title)
                    title_run.bold = True
                    title_run.font.size = shared.Pt(14)

                    for line in chapter_text.splitlines():
                        self.doc.add_paragraph(line)

                    self.doc.add_page_break()

                self.doc.save(self.docx_file_path)
                self.chapter_buffer.clear()

            except Exception as e:
                self.msg_manager.show_message('there_was_an_error', e)

    def save_chapter(self, chapter_title, header_p):

        try:
            all_text = '\n'.join([p.get_text(strip=True) for p in header_p if p.get_text(strip=True)])
            cleaned_text = '\n'.join([line.strip() for line in all_text.splitlines() if line.strip()])

            self.chapter_buffer.append((chapter_title, cleaned_text))

            if len(self.chapter_buffer) >= self.buffer_limit:
                self.flush_buffer()

        except Exception as e:
            self.msg_manager.show_message('there_was_an_error', e)

    def finalize(self):
        if self.chapter_buffer:
            self.flush_buffer()
